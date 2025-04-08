import os 
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from PIL import Image
import io

# 配置路径
pdf_folder = '待处理文档' # 待处理文档文件夹路径
output_folder = os.path.join(pdf_folder, pdf_folder + '_'+'docx')
os.makedirs(output_folder, exist_ok=True)

# 存放生成word文件的路径
word_files = []

# 最大宽高 (A4页面约为6x9英寸，缩小90%)
MAX_WIDTH_INCH = 6.0 * 0.9
MAX_HEIGHT_INCH = 9.0 * 0.9

# 缩放图片90%
def resize_image(pil_img):
    width, height = pil_img.size
    aspect_ratio = height / width
    new_width_px = int(MAX_WIDTH_INCH * 300)  # DPI=300
    new_height_px = int(new_width_px * aspect_ratio)

    # 超过高度则限制高度
    if new_height_px > MAX_HEIGHT_INCH * 300:
        new_height_px = int(MAX_HEIGHT_INCH * 300)
        new_width_px = int(new_height_px / aspect_ratio)

    return pil_img.resize((new_width_px, new_height_px), Image.LANCZOS)

# 遍历所有PDF文件
for pdf_file in os.listdir(pdf_folder):
    if not pdf_file.lower().endswith('.pdf'):
        continue

    pdf_path = os.path.join(pdf_folder, pdf_file)
    doc = fitz.open(pdf_path)
    word_doc = Document()

    page_added = False
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(dpi=200)
        
        pil_img = Image.open(io.BytesIO(pix.tobytes('png')))
        pil_img = resize_image(pil_img)  # 统一缩放到90%
        
        img_byte_arr = io.BytesIO()
        pil_img.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)

        if page_added:
            word_doc.add_page_break()

        #word_doc.add_paragraph(pdf_file)
        word_doc.add_paragraph(os.path.splitext(pdf_file)[0],style='Heading 1')

        word_doc.add_picture(img_byte_arr, width=Inches(MAX_WIDTH_INCH))
        page_added = True

    if page_added:
        word_file_path = os.path.join(output_folder, pdf_file.replace('.pdf', '.docx'))
        word_doc.save(word_file_path)
        word_files.append(word_file_path)
    else:
        print(f"{pdf_file} 未插入任何有效页面")

print('所有文件处理完成，无空白页！')
