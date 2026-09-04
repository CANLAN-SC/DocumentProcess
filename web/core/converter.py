import io
from dataclasses import dataclass
from pathlib import Path

import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docxcompose.composer import Composer
from PIL import Image, ImageOps


A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
HEADING_SPACE_CM = 1.35
SUPPORTED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png"}


@dataclass(frozen=True)
class LayoutConfig:
    """一批 Word 文档共用的 A4 页面与图片版式。"""

    image_scale_percent: float = 90.0
    margin_top_cm: float = 2.54
    margin_bottom_cm: float = 2.54
    margin_left_cm: float = 2.54
    margin_right_cm: float = 2.54

    @property
    def available_width_cm(self):
        return A4_WIDTH_CM - self.margin_left_cm - self.margin_right_cm

    @property
    def available_image_height_cm(self):
        return A4_HEIGHT_CM - self.margin_top_cm - self.margin_bottom_cm - HEADING_SPACE_CM

    def validate(self):
        margins = (
            self.margin_top_cm,
            self.margin_bottom_cm,
            self.margin_left_cm,
            self.margin_right_cm,
        )
        if any(value < 0 or value > 8 for value in margins):
            raise ValueError("页边距必须在 0 到 8 厘米之间")
        if not 10 <= self.image_scale_percent <= 100:
            raise ValueError("图片大小必须在 10% 到 100% 之间")
        if self.available_width_cm <= 1:
            raise ValueError("左右页边距之和过大，页面可用宽度不足")
        if self.available_image_height_cm <= 1:
            raise ValueError("上下页边距之和过大，页面可用高度不足")

    def image_size_cm(self, pixel_width, pixel_height):
        if pixel_width <= 0 or pixel_height <= 0:
            raise ValueError("图片尺寸无效")
        scale = self.image_scale_percent / 100.0
        max_width = self.available_width_cm * scale
        max_height = self.available_image_height_cm * scale
        ratio = pixel_height / pixel_width
        width = max_width
        height = width * ratio
        if height > max_height:
            height = max_height
            width = height / ratio
        return width, height


class WordConverter:
    """使用 PyMuPDF 渲染 PDF，不依赖 Poppler 或系统 PATH。"""

    def __init__(self, layout=None, compress=False):
        self.layout = layout or LayoutConfig()
        self.layout.validate()
        self.compress = compress

    @staticmethod
    def _set_font(document, font_name="微软雅黑", font_size=12):
        style = document.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _configure_document(self, document):
        section = document.sections[0]
        section.page_width = Cm(A4_WIDTH_CM)
        section.page_height = Cm(A4_HEIGHT_CM)
        section.top_margin = Cm(self.layout.margin_top_cm)
        section.bottom_margin = Cm(self.layout.margin_bottom_cm)
        section.left_margin = Cm(self.layout.margin_left_cm)
        section.right_margin = Cm(self.layout.margin_right_cm)

    @staticmethod
    def _add_heading(document, title, page_break_before=False):
        heading = document.add_paragraph(style="Heading 1")
        heading.paragraph_format.page_break_before = page_break_before
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run(title)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "微软雅黑"
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def _prepare_image(self, pil_image, width_cm, height_cm):
        target_width_px = max(1, int(width_cm / 2.54 * 300))
        target_height_px = max(1, int(height_cm / 2.54 * 300))
        image = ImageOps.exif_transpose(pil_image).copy()
        image.thumbnail((target_width_px, target_height_px), Image.Resampling.LANCZOS)
        buffer = io.BytesIO()
        if self.compress:
            if image.mode != "RGB":
                background = Image.new("RGB", image.size, "white")
                if "A" in image.getbands():
                    background.paste(image, mask=image.getchannel("A"))
                else:
                    background.paste(image)
                image = background
            image.save(buffer, format="JPEG", quality=80, optimize=True)
        else:
            if image.mode not in ("1", "L", "LA", "P", "RGB", "RGBA"):
                image = image.convert("RGB")
            image.save(buffer, format="PNG", optimize=True)
        buffer.seek(0)
        return buffer

    def _add_image_page(self, document, pil_image, title, page_break_before=False):
        self._add_heading(document, title, page_break_before)
        width_cm, height_cm = self.layout.image_size_cm(*pil_image.size)
        image_buffer = self._prepare_image(pil_image, width_cm, height_cm)
        picture_paragraph = document.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.space_before = Pt(0)
        picture_paragraph.paragraph_format.space_after = Pt(0)
        picture_paragraph.add_run().add_picture(
            image_buffer, width=Cm(width_cm), height=Cm(height_cm)
        )

    def _new_document(self):
        document = Document()
        self._set_font(document)
        self._configure_document(document)
        return document

    @staticmethod
    def _pdf_page_image(page, dpi):
        scale = dpi / 72.0
        pixmap = page.get_pixmap(matrix=pymupdf.Matrix(scale, scale), alpha=False)
        with Image.open(io.BytesIO(pixmap.tobytes("png"))) as image:
            image.load()
            return image.convert("RGB")

    def convert_pdf(self, input_path, output_path):
        title = Path(input_path).stem
        document = self._new_document()
        dpi = 200 if self.compress else 260
        with pymupdf.open(input_path) as pdf:
            if pdf.page_count == 0:
                raise ValueError("PDF 中没有有效页面")
            for index, page in enumerate(pdf):
                image = self._pdf_page_image(page, dpi)
                self._add_image_page(document, image, title, index > 0)
                image.close()
        document.save(output_path)
        return output_path

    def convert_image(self, input_path, output_path):
        title = Path(input_path).stem
        document = self._new_document()
        with Image.open(input_path) as image:
            image.load()
            self._add_image_page(document, image, title)
        document.save(output_path)
        return output_path

    def convert(self, input_path, output_path):
        extension = Path(input_path).suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件类型：{extension}")
        if extension == ".pdf":
            return self.convert_pdf(input_path, output_path)
        return self.convert_image(input_path, output_path)

    @staticmethod
    def merge_documents(word_files, output_path):
        if not word_files:
            raise ValueError("没有可合并的 Word 文档")
        master = Document(word_files[0])
        composer = Composer(master)
        for word_file in word_files[1:]:
            composer.append(Document(word_file))
        composer.save(output_path)
        return output_path


def render_preview(input_path, output_folder, page_limit=6, dpi=110):
    """生成浏览器预览图，返回（页面信息列表，总页数）。"""
    input_path = Path(input_path)
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)
    extension = input_path.suffix.lower()
    pages = []

    if extension == ".pdf":
        with pymupdf.open(input_path) as pdf:
            total_pages = pdf.page_count
            if total_pages == 0:
                raise ValueError("PDF 中没有有效页面")
            for index in range(min(total_pages, page_limit)):
                image = WordConverter._pdf_page_image(pdf[index], dpi)
                preview_path = output_folder / f"page-{index + 1}.jpg"
                image.save(preview_path, "JPEG", quality=86, optimize=True)
                pages.append(
                    {
                        "filename": preview_path.name,
                        "width": image.width,
                        "height": image.height,
                        "page": index + 1,
                    }
                )
                image.close()
            return pages, total_pages

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型：{extension}")
    with Image.open(input_path) as image:
        image = ImageOps.exif_transpose(image).convert("RGB")
        image.thumbnail((1800, 2400), Image.Resampling.LANCZOS)
        preview_path = output_folder / "page-1.jpg"
        image.save(preview_path, "JPEG", quality=88, optimize=True)
        pages.append(
            {
                "filename": preview_path.name,
                "width": image.width,
                "height": image.height,
                "page": 1,
            }
        )
    return pages, 1
