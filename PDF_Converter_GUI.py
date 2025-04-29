import os
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from PIL import Image
import io
from pdf2image import convert_from_path
from docxcompose.composer import Composer
from docx.shared import RGBColor  

class PDFConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF/图片转Word工具 v2.0")
        self.input_folder = ""
        self.compress_mode = False
        self.merge_files = False
        if getattr(sys, 'frozen', False):
            # cx_Freeze的路径基地址是sys.executable的目录
            base_path = os.path.dirname(sys.executable)
            self.poppler_path = os.path.join(base_path, "lib", "poppler", "Library", "bin")
        else:
            self.poppler_path = os.path.join("poppler", "Library", "bin")


        # GUI布局
        ttk.Label(root, text="输入文件夹:").grid(row=0, column=0, padx=5, pady=5)
        self.entry_path = ttk.Entry(root, width=40)
        self.entry_path.grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(root, text="浏览", command=self.select_folder).grid(row=0, column=2, padx=5, pady=5)

        # 模式选择
        self.mode_var = tk.StringVar(value="normal")
        ttk.Radiobutton(root, text="高清模式", variable=self.mode_var, value="normal").grid(row=1, column=1, sticky=tk.W)
        ttk.Radiobutton(root, text="压缩模式", variable=self.mode_var, value="compress").grid(row=2, column=1, sticky=tk.W)
        
        # 合并选项
        self.merge_var = tk.BooleanVar()
        ttk.Checkbutton(root, text="自动合并文档", variable=self.merge_var).grid(row=3, column=1, sticky=tk.W)
        
        # 操作按钮
        ttk.Button(root, text="开始转换", command=self.start_conversion).grid(row=4, column=1, pady=10)
        
        # 日志区域
        self.log_text = tk.Text(root, height=10, width=50)
        self.log_text.grid(row=5, column=0, columnspan=3, padx=10, pady=10)

        # 全局配置
        self.MAX_WIDTH_INCH = 6.0 * 0.9
        self.MAX_HEIGHT_INCH = 9.0 * 0.9

    def select_folder(self):
        """选择输入文件夹"""
        self.input_folder = filedialog.askdirectory()
        self.entry_path.delete(0, tk.END)
        self.entry_path.insert(0, self.input_folder)
        self.log(f"已选择文件夹: {self.input_folder}")

    def log(self, message):
        """记录日志"""
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def set_font(self, document, font_name='微软雅黑', font_size=12):
        """统一设置字体"""
        style = document.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def resize_image(self, pil_img):
        """图片缩放逻辑"""
        width, height = pil_img.size
        aspect_ratio = height / width
        new_width_px = int(self.MAX_WIDTH_INCH * 300)
        new_height_px = int(new_width_px * aspect_ratio)

        if new_height_px > self.MAX_HEIGHT_INCH * 300:
            new_height_px = int(self.MAX_HEIGHT_INCH * 300)
            new_width_px = int(new_height_px / aspect_ratio)

        return pil_img.resize((new_width_px, new_height_px), Image.LANCZOS)

    def process_pdf(self, file_path, output_folder, compress=False):
        """处理PDF文件"""
        try:
            pages = convert_from_path(file_path, dpi=200, poppler_path=self.poppler_path)
            word_doc = Document()
            self.set_font(word_doc)

            page_added = False
            for page in pages:
                pil_img = self.resize_image(page)
                
                if compress:
                    img_buffer = io.BytesIO()
                    pil_img.save(img_buffer, format='JPEG', quality=80, optimize=True)
                else:
                    img_buffer = io.BytesIO()
                    pil_img.save(img_buffer, format='PNG')

                img_buffer.seek(0)

                if page_added:
                    word_doc.add_page_break()

                # 添加标题
                heading = word_doc.add_paragraph(os.path.splitext(os.path.basename(file_path))[0], style='Heading 1')
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(0, 0, 0)  # 添加这一行确保标题也是黑色

                run.font.name = '微软雅黑'
                run.font.size = Pt(14)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

                word_doc.add_picture(img_buffer, width=Inches(self.MAX_WIDTH_INCH))
                page_added = True

            if page_added:
                output_path = os.path.join(output_folder, os.path.basename(file_path).replace('.pdf', '.docx'))
                word_doc.save(output_path)
                self.log(f"转换成功: {os.path.basename(file_path)}")
            return True
        except Exception as e:
            self.log(f"转换失败 {file_path}: {str(e)}")
            return False

    def process_image(self, file_path, output_folder):
        """处理图片文件"""
        try:
            pil_img = Image.open(file_path)
            pil_img = self.resize_image(pil_img)
            word_doc = Document()
            self.set_font(word_doc)

            img_buffer = io.BytesIO()
            pil_img.save(img_buffer, format='PNG')
            img_buffer.seek(0)

            # 添加标题
            heading = word_doc.add_paragraph(os.path.splitext(os.path.basename(file_path))[0], style='Heading 1')
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)  # 添加这一行确保标题也是黑色

            run.font.name = '微软雅黑'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            word_doc.add_picture(img_buffer, width=Inches(self.MAX_WIDTH_INCH))
            
            output_path = os.path.join(output_folder, os.path.splitext(os.path.basename(file_path))[0] + '.docx')
            word_doc.save(output_path)
            self.log(f"转换成功: {os.path.basename(file_path)}")
            return True
        except Exception as e:
            self.log(f"转换失败 {file_path}: {str(e)}")
            return False

    def merge_documents(self, input_folder):
        """合并Word文档"""
        try:
            word_folder = os.path.join(input_folder, os.path.basename(input_folder) + '_docx')
            output_path = os.path.join(input_folder, os.path.basename(input_folder) + '_合并.docx')
            
            docx_files = sorted(
                [f for f in os.listdir(word_folder) if f.lower().endswith('.docx')],
                key=lambda x: x.lower()
            )

            if len(docx_files) == 0:
                self.log("没有找到可合并的Word文件")
                return

            master = Document(os.path.join(word_folder, docx_files[0]))
            composer = Composer(master)

            for docx_file in docx_files[1:]:
                sub_doc = Document(os.path.join(word_folder, docx_file))
                composer.append(sub_doc)

            composer.save(output_path)
            self.log(f"合并完成: {output_path}")
        except Exception as e:
            self.log(f"合并失败: {str(e)}")

    def start_conversion(self):
        """启动转换线程"""
        if not self.input_folder:
            messagebox.showerror("错误", "请先选择输入文件夹")
            return

        compress_mode = self.mode_var.get() == "compress"
        need_merge = self.merge_var.get()
        output_folder = os.path.join(self.input_folder, os.path.basename(self.input_folder) + '_docx')
        os.makedirs(output_folder, exist_ok=True)

        def conversion_thread():
            try:
                # 处理所有文件
                for file in os.listdir(self.input_folder):
                    file_path = os.path.join(self.input_folder, file)
                    if file.lower().endswith('.pdf'):
                        self.process_pdf(file_path, output_folder, compress=compress_mode)
                    elif file.lower().endswith(('.jpg', '.jpeg', '.png')):
                        self.process_image(file_path, output_folder)
                
                # 合并文档
                if need_merge:
                    self.merge_documents(self.input_folder)
                
                self.log("所有操作已完成！")
                messagebox.showinfo("完成", "所有操作已完成！")
            except Exception as e:
                self.log(f"发生严重错误: {str(e)}")
                messagebox.showerror("错误", str(e))

        threading.Thread(target=conversion_thread, daemon=True).start()

if __name__ == "__main__":
    root = tk.Tk()
    app = PDFConverterApp(root)
    root.mainloop()