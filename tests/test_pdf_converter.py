import io
import os
import tempfile
import unittest
import zipfile

import pymupdf
from docx import Document
from PIL import Image

from web.app import create_app
from web.core.converter import LayoutConfig, WordConverter, render_preview


def png_bytes(size=(1200, 1800), color="white"):
    buffer = io.BytesIO()
    Image.new("RGB", size, color).save(buffer, "PNG")
    buffer.seek(0)
    return buffer


class LayoutConfigTests(unittest.TestCase):
    def test_portrait_image_fits_scaled_printable_area(self):
        config = LayoutConfig(
            image_scale_percent=75,
            margin_top_cm=1.5,
            margin_bottom_cm=1.6,
            margin_left_cm=2.0,
            margin_right_cm=2.1,
        )
        width, height = config.image_size_cm(1200, 1800)
        self.assertAlmostEqual(width, 12.625, places=3)
        self.assertAlmostEqual(height, 18.9375, places=3)

    def test_invalid_margins_are_rejected(self):
        with self.assertRaises(ValueError):
            LayoutConfig(margin_left_cm=8.1, margin_right_cm=2).validate()

    def test_image_scale_200_is_accepted(self):
        LayoutConfig(image_scale_percent=200).validate()

    def test_image_scale_above_200_is_rejected(self):
        with self.assertRaises(ValueError) as ctx:
            LayoutConfig(image_scale_percent=201).validate()
        self.assertIn("200", str(ctx.exception))

    def test_image_size_doubles_at_200_percent(self):
        base = LayoutConfig(image_scale_percent=100)
        doubled = LayoutConfig(image_scale_percent=200)
        # 竖图触发「高度超限按比例回缩」，横图触发「仅宽度约束」两条分支，
        # 两种情况下 200% 的宽高都应恰好是 100% 的 2 倍。
        for size in ((1200, 1800), (1800, 1200)):
            base_w, base_h = base.image_size_cm(*size)
            double_w, double_h = doubled.image_size_cm(*size)
            self.assertAlmostEqual(double_w, base_w * 2, places=3)
            self.assertAlmostEqual(double_h, base_h * 2, places=3)


class DocumentOutputTests(unittest.TestCase):
    def test_image_size_and_margins_are_written_to_docx(self):
        config = LayoutConfig(
            image_scale_percent=75,
            margin_top_cm=1.5,
            margin_bottom_cm=1.6,
            margin_left_cm=2.0,
            margin_right_cm=2.1,
        )
        with tempfile.TemporaryDirectory() as folder:
            image_path = os.path.join(folder, "sample.png")
            output_path = os.path.join(folder, "sample.docx")
            Image.new("RGB", (1200, 1800), "white").save(image_path)
            WordConverter(config).convert(image_path, output_path)
            document = Document(output_path)

        section = document.sections[0]
        picture = document.inline_shapes[0]
        expected_width, expected_height = config.image_size_cm(1200, 1800)
        self.assertAlmostEqual(section.top_margin.cm, 1.5, places=2)
        self.assertAlmostEqual(section.bottom_margin.cm, 1.6, places=2)
        self.assertAlmostEqual(section.left_margin.cm, 2.0, places=2)
        self.assertAlmostEqual(section.right_margin.cm, 2.1, places=2)
        self.assertAlmostEqual(picture.width.cm, expected_width, places=2)
        self.assertAlmostEqual(picture.height.cm, expected_height, places=2)

    def test_pdf_conversion_and_preview_do_not_require_poppler(self):
        with tempfile.TemporaryDirectory() as folder:
            pdf_path = os.path.join(folder, "two-pages.pdf")
            output_path = os.path.join(folder, "two-pages.docx")
            preview_folder = os.path.join(folder, "preview")
            pdf = pymupdf.open()
            pdf.new_page().insert_text((72, 72), "Page one")
            pdf.new_page().insert_text((72, 72), "Page two")
            pdf.save(pdf_path)
            pdf.close()

            pages, total_pages = render_preview(pdf_path, preview_folder)
            WordConverter().convert(pdf_path, output_path)
            document = Document(output_path)

        self.assertEqual(total_pages, 2)
        self.assertEqual(len(pages), 2)
        self.assertEqual(len(document.inline_shapes), 2)
        self.assertTrue(document.paragraphs[2].paragraph_format.page_break_before)


class WebApiTests(unittest.TestCase):
    def setUp(self):
        self.temp_folder = tempfile.TemporaryDirectory()
        self.app = create_app(
            {
                "TESTING": True,
                "DATA_DIR": self.temp_folder.name,
                "MAX_CONTENT_LENGTH": 20 * 1024 * 1024,
            }
        )
        self.client = self.app.test_client()

    def tearDown(self):
        self.temp_folder.cleanup()

    def test_index_loads_unified_web_interface(self):
        response = self.client.get("/")
        self.assertEqual(response.status_code, 200)
        self.assertIn("Word 实时预览".encode(), response.data)

    def test_preview_and_convert_image(self):
        preview_response = self.client.post(
            "/api/preview",
            data={"file": (png_bytes(), "示例.png")},
            content_type="multipart/form-data",
        )
        self.assertEqual(preview_response.status_code, 200)
        preview_data = preview_response.get_json()
        self.assertEqual(preview_data["total_pages"], 1)
        preview_asset = self.client.get(preview_data["pages"][0]["url"])
        self.assertEqual(preview_asset.status_code, 200)
        preview_asset.close()

        convert_response = self.client.post(
            "/api/convert",
            data={
                "files": (png_bytes(), "示例.png"),
                "image_scale": "70",
                "margin_top": "1.2",
                "margin_bottom": "1.3",
                "margin_left": "1.4",
                "margin_right": "1.5",
                "quality": "normal",
                "merge": "false",
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(convert_response.status_code, 200)
        result = convert_response.get_json()
        download = self.client.get(result["download_url"])
        self.assertEqual(download.status_code, 200)
        document = Document(io.BytesIO(download.data))
        self.assertAlmostEqual(document.sections[0].top_margin.cm, 1.2, places=2)
        self.assertEqual(len(document.inline_shapes), 1)
        download.close()

    def test_invalid_pdf_returns_readable_error_without_poppler_message(self):
        response = self.client.post(
            "/api/preview",
            data={"file": (io.BytesIO(b"not a pdf"), "broken.pdf")},
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 400)
        self.assertNotIn("poppler", response.get_json()["error"].lower())

    def test_multiple_files_with_same_stem_are_packaged_without_overwrite(self):
        response = self.client.post(
            "/api/convert",
            data={
                "files": [
                    (png_bytes(color="white"), "same.png"),
                    (png_bytes(color="#eeeeee"), "same.jpg"),
                ],
                "merge": "false",
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 200)
        download = self.client.get(response.get_json()["download_url"])
        with zipfile.ZipFile(io.BytesIO(download.data)) as archive:
            self.assertEqual(archive.namelist(), ["same.docx", "same_2.docx"])
        download.close()


if __name__ == "__main__":
    unittest.main()
