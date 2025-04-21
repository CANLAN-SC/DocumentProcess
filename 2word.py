import os
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
from pdf2image import convert_from_path

# 配置路径
input_folder = '专利'  # 待处理文件夹路径
output_folder = os.path.join(input_folder, input_folder + '_' + 'docx')
os.makedirs(output_folder, exist_ok=True)

poppler_path = r"poppler\Library\bin"  # 你的poppler路径，Windows用户不需要修改

# 最大宽高 (A4页面约为6x9英寸，缩小90%)
MAX_WIDTH_INCH = 6.0 * 0.85
MAX_HEIGHT_INCH = 9.0 * 0.85

# 缩放图片90%
def resize_image(pil_img):
    width, height = pil_img.size
    aspect_ratio = height / width
    new_width_px = int(MAX_WIDTH_INCH * 300)  # DPI=300
    new_height_px = int(new_width_px * aspect_ratio)

    if new_height_px > MAX_HEIGHT_INCH * 300:
        new_height_px = int(MAX_HEIGHT_INCH * 300)
        new_width_px = int(new_height_px / aspect_ratio)

    return pil_img.resize((new_width_px, new_height_px), Image.LANCZOS)

# 处理文件
for file in os.listdir(input_folder):
    file_path = os.path.join(input_folder, file)
    file_lower = file.lower()
    
    word_doc = Document()

    if file_lower.endswith('.pdf'):
        pages = convert_from_path(file_path, dpi=200, poppler_path=poppler_path)

        #pages = convert_from_path(file_path, dpi=200)

        page_added = False
        for page in pages:
            pil_img = resize_image(page)

            img_byte_arr = io.BytesIO()
            pil_img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)

            if page_added:
                word_doc.add_page_break()

            word_doc.add_paragraph(os.path.splitext(file)[0], style='Heading 1')
            word_doc.add_picture(img_byte_arr, width=Inches(MAX_WIDTH_INCH))
            page_added = True

        if page_added:
            word_file_path = os.path.join(output_folder, file.replace('.pdf', '.docx'))
            word_doc.save(word_file_path)
            print(f"PDF {file} 转换成功")
        else:
            print(f"PDF {file} 未插入任何有效页面")

    elif file_lower.endswith(('.jpg', '.jpeg', '.png')):
        pil_img = Image.open(file_path)
        pil_img = resize_image(pil_img)

        img_byte_arr = io.BytesIO()
        pil_img.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)

        word_doc.add_paragraph(os.path.splitext(file)[0], style='Heading 1')
        word_doc.add_picture(img_byte_arr, width=Inches(MAX_WIDTH_INCH))

        word_file_path = os.path.join(output_folder, os.path.splitext(file)[0] + '.docx')
        word_doc.save(word_file_path)
        print(f"图片 {file} 转换成功")

print('所有文件处理完成，无空白页！')
